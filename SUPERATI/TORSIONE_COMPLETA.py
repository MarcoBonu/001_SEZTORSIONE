#TORSIONE COMPLETA
import numpy as np
import math 
import matplotlib.pyplot as plt
import docx
from matplotlib.patches import Rectangle
from docx.shared import Inches
import pandas as pd
np.set_printoptions(suppress=True,linewidth=np.nan)

def omega(k,zi):
    lista=[np.cosh(k*zi), np.sinh(k*zi), zi, 1]
    omega=np.array(lista)
    return omega

def omega_p(k,zi):
    lista=[k*np.sinh(k*zi), k*np.cosh(k*zi), 1, 0]
    omega_p=np.array(lista)
    return omega_p

def omega_pp(k,zi):
    lista=[k**2*np.cosh(k*zi), k**2*np.sinh(k*zi), 0, 0]
    omega_pp=np.array(lista)
    return omega_pp

def omega_ppp(k,zi):
    lista=[k**3*np.sinh(k*zi), k**3*np.cosh(k*zi), 0, 0]
    omega_ppp=np.array(lista)
    return omega_ppp

def kj(Lj,Jt,J_psi,E,G):
    kj=Lj*math.sqrt((G*Jt)/(E*J_psi))
    return kj

def materiale(flag):

    if 1 == flag:
        mat = "acciaio"
    else:
        mat = "calcestruzzo"

    dict_materiali={"acciaio": [210000, 80000],
                    "calcestruzzo" : [30000 , 15000]}
    
    E_j = dict_materiali[mat][0]
    G_j = dict_materiali[mat][1]

    return [E_j, G_j,mat]

def Inerzie_torsionali(I):

    J_t=[]
    J_psi=[]

    flag_trave_prismatica=0
    while ("S" != flag_trave_prismatica and "N" != flag_trave_prismatica):
        flag_trave_prismatica = input("La trave è prismatica? (Stessa sezione lungo l'asse baricentrico) [S/N] : ")

    if "S"==flag_trave_prismatica:
        Jt_j=float(input("Indicare Jt della trave [cm^4] : "))
        Jpsi_j=float(input("Indicare J_psi della trave [cm^6] : "))
        for j in range(I+1):
            J_t.append(Jt_j)
            J_psi.append(Jpsi_j)

    else:

        for j in range(I+1):
            Jt_j = float(input(f"Indicare J_t del tratto {j+1} [cm^4] : "))
            J_t.append(Jt_j) 

        for j in range(I+1):
            Jpsi_j=float(input(f"Indicare J_psi del tratto {j+1} [cm^4] : "))
            J_psi.append(Jpsi_j)

    return [J_t , J_psi]

def cc_incastro_dsv(j,I,K,Lj,flag):
    kj = K[j]
    zj_sx = 0
    zj_dx = Lj[j]/Lj[j]


    #flag = 1 resistuisce la cc dell'incastro esterno flag2 restituisce la cc della continuità
    if j==0:
        if flag == 2:
            kp = K[j+1]
            zp = 0
            cc_incastro_dsv = [ *omega(kj,zj_dx), *-omega(kp,zp)]
        elif flag== 1:
            cc_incastro_dsv = [*omega(kj,zj_sx)]

    elif j==I:
        cc_incastro_dsv = [*omega(kj,zj_dx)]
    else:
        kp = K[j+1]
        zp = 0
        cc_incastro_dsv = [ *omega(kj,zj_dx), *-omega(kp,zp)]
    #print("Incastro dsv")
    #print(cc_incastro_dsv)
    #input("AVANTI?")
    return cc_incastro_dsv

def cc_incastro_bimom(j,I,K,Lj,flag):
    kj = K[j]
    zj_sx = 0
    zj_dx = Lj[j]/Lj[j]

    #flag = 1 resistuisce la cc dell'incastro esterno flag2 restituisce la cc della continuità
    if j == 0 :
        if flag == 2:
            kp = K[j+1]
            zp = 0
            cc_incastro_bimom = [*omega_p(kj,zj_dx) , *-omega_p(kp,zp)]
            
        elif flag==1:
            cc_incastro_bimom = [*omega_p(kj,zj_sx)]
    elif j==I:
        cc_incastro_bimom = [*omega_p(kj,zj_dx)]
    else:
        kp = K[j+1]
        zp = 0
        cc_incastro_bimom = [*omega_p(kj,zj_dx) , *-omega_p(kp,zp)]
        
    #print("Incastro bimomento")
    #print(cc_incastro_bimom)
    #input("AVANTI?")
    return cc_incastro_bimom

def cc_equ_mom(j,I,K,Lj,J_t, J_psi, E,G, M3_i):
    kj = K[j]
    zj_sx=0
    zj_dx = Lj[j]/Lj[j]
    Jt_j = J_t[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    G_j = G[j]

    kp = K[j+1]
    zp_dx=Lj[j+1]/Lj[j+1]
    zp_sx =0
    zp = 0
    Jt_p = J_t[j+1]
    Jpsi_p = J_psi[j+1]
    E_p = E[j+1]
    G_p = G[j+1]

    MT_j = G_j*Jt_j* (omega_p(kj,zj_dx))/Lj[j] - E_j * Jpsi_j*(omega_ppp(kj,zj_dx))/Lj[j]**3
    MT_p = G_p*Jt_p* (omega_p(kp,zp_sx))/Lj[j+1] - E_p * Jpsi_p*(omega_ppp(kp,zp_sx))/Lj[j+1]**3
    #print("Equ Momento torcente")
    #print([ *MT_j , *-MT_p])
    #input("AVANTI?")
    return [ *MT_j , *-MT_p]

def cc_equ_bimom(j,K,Lj, J_psi, E):
    kj = K[j]
    zj = Lj[j]/Lj[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    
    kp = K[j+1]
    zp = 0
    Jpsi_p = J_psi[j+1]
    E_p = E[j+1]
    

    B_j = E_j * Jpsi_j*omega_pp(kj,zj)/Lj[j]**2
    B_p = -E_p * Jpsi_p*omega_pp(kp,zp)/Lj[j+1]**2
    #print("Equ Bimomento")
    #print([*B_j ,*-B_p ])
    #input("AVANTI?")
    return [*B_j ,*B_p ]

def cc_lembo_libero_M(j,K,Lj,J_t, J_psi, E,G):
    kj = K[j]
    zj_sx=0
    zj_dx = Lj[j]/Lj[j]
    Jt_j = J_t[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    G_j = G[j]

    MT_j = G_j*Jt_j* (omega_p(kj,zj_dx))/Lj[j] - E_j * Jpsi_j*(omega_ppp(kj,zj_dx))/Lj[j]**3

    return MT_j

def cc_lembo_libero_B(j,K,Lj, J_psi, E):
    kj = K[j]
    zj = Lj[j]/Lj[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    
    B_j = E_j * Jpsi_j*omega_pp(kj,zj)/Lj[j]**2

    return B_j 

def Mt(z,j,K,Lj,J_t, J_psi, E,G , X):  
    
    kj = K[j]
    zj= z/Lj[j]
    Jt_j = J_t[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    G_j = G[j]
    Xj = X[4*j:4*j+4]

    MT_P = G_j*Jt_j* np.matmul(omega_p(kj,zj),Xj)/Lj[j]
    MT_S = - E_j * Jpsi_j*np.matmul(omega_ppp(kj,zj),Xj)/Lj[j]**3
    MT =  MT_P + MT_S

    return  MT_P, MT_S, MT 

def B(j,z,K,Lj, J_psi, E, X): 
    kj = K[j]
    zj = z/Lj[j]
    Jpsi_j = J_psi[j]
    E_j = E[j]
    Xj = X[4*j:4*j+4]
    
    B = E_j * Jpsi_j*np.matmul(omega_pp(kj,zj),Xj)/Lj[j]**2

    return  B

def W(zj,j,K ,Lj, X):
    
    kj = K[j]
    zj=zj/Lj[j]
    Xj = X[4*j:4*j+4]

    W = np.matmul(omega(kj,zj),Xj)
    
    return W

def Wp(zj,j,K ,Lj, X):
    
    kj = K[j]
    zj=zj/Lj[j]
    Xj = X[4*j:4*j+4]

    Wp = np.matmul(omega_p(kj,zj),Xj)
    
    return Wp

def grafici( I,K,Lj,J_t, J_psi, E,G , X,Z,flag_lembo_vincolato):
    
    BIM=[]
    OMEGA=[]
    MTTOT=[]
    MTSEC=[]
    MTPRIM=[]
    OMEGA_P=[]
    p=1
   

    for j in range (I+1):
        
        z=range(0,int(Lj[j]),p)
        for zi in z:
            BIM.append(B(j,zi,K,Lj, J_psi, E,X))
            OMEGA_P.append(Wp(zi,j,K ,Lj, X))
            MTTOT.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[2][0])
            MTSEC.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[1][0])
            MTPRIM.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[0][0])
            OMEGA.append(W(zi,j,K ,Lj, X))

    z=[]
    for j in range (I+1):
        if j==0:
            start = int(0)
        else:
            start  = z[-1]+p
        
        zj=range(start,start+int(Lj[j]),p)
        
        z.extend(zj)
    
    plt.subplot(311)
    plt.plot(z, OMEGA,label="Rotazione w")
    plt.legend(loc="upper left")
    plt.grid()
    plt.xlabel("z  [cm]")

    plt.subplot(312)
    plt.plot(z, OMEGA_P,label="Derivata prima della rotazione")
    plt.legend(loc="upper left")
    plt.grid()
    plt.xlabel("z  [cm]")

    plt.subplot(313)
    plt.plot(z, BIM, label="Bimomento")
    plt.legend(loc="upper left")
    plt.grid()
    plt.xlabel("z  [cm]")

    plt.savefig("ALTRI GRAFICI.png",bbox_inches='tight')
    
    fig, ax = plt.subplots()
    ax.plot(z, MTTOT,"r", label="Momento torcente totale")
    ax.plot(z, MTPRIM,"b", label="Momento torcente primario")
    ax.plot(z, MTSEC,"g", label="Momento torcente secondario")
    ax.plot(z,np.zeros(len(z)),"k",linewidth=5)
    largh=int(90)
    altez=abs(max(MTTOT, key=abs))*2
    ax.add_patch(Rectangle([-largh,-altez/2], largh, altez, angle=0.0,facecolor = 'white', ec="black",hatch = '//'))
    if flag_lembo_vincolato=="S":
        ax.add_patch(Rectangle([int(z[-1]),-altez/2], largh, altez, angle=0.0,facecolor = 'white', ec="black",hatch = '//'))
    for j in range(I):
        zi = Z[j+1]
        largh_rett_rosso=4
        ax.add_patch(Rectangle([zi-largh_rett_rosso/2,-altez/8],largh_rett_rosso, altez/4, angle=0.0,facecolor = 'black'))

    ax = plt.gca()
    #ax.set_aspect('equal', adjustable='box')
    ax.legend(loc="upper center",bbox_to_anchor=(0.5, -0.15),ncols=3)
    plt.xlabel("z [cm]")
    plt.ylabel("[kNcm]")
    plt.grid()
    plt.savefig("MOMENTO TORCENTE.png",bbox_inches='tight') 

def plot_geometria( I,Lj,Z,flag_lembo_libero):
    z=[]
    p=1
    for j in range (I+1):
        if j==0:
            start = int(0)
        else:
            start  = z[-1]+p
        
        zj=range(start,start+int(Lj[j]),p)
        
        z.extend(zj)


    fig, ax = plt.subplots()
    ax.plot(z,np.zeros(len(z)),"k")
    largh=int(30)
    altez=int(90)
    ax.add_patch(Rectangle([-largh,-altez/2], largh, altez, angle=0.0,facecolor = 'white', ec="black",hatch = '//'))
    if flag_lembo_libero=="S":
        ax.add_patch(Rectangle([int(z[-1]),-altez/2], largh, altez, angle=0.0,facecolor = 'white', ec="black",hatch = '//'))

    for j in range(I):
        zi = Z[j+1]
        ax.add_patch(Rectangle([zi-2.5,-altez/8], 2, altez/4, angle=0.0,facecolor = 'red'))

        largh_freccia=int(5)
        spost_freccia=int(8)
        plt.arrow(zi,altez/4+10,10,0,ec="red",fc="red",width=largh_freccia)
        plt.arrow(zi+spost_freccia,altez/4+10,10,0,ec="red",fc="red", width=largh_freccia)
        ax.text(zi-2*spost_freccia, altez/4+30, f'Mt{j+1}',color="red")
    
    
    
    ax = plt.gca()
    ax.set_aspect('equal', adjustable='box')
    plt.xlabel("z [cm]")
    plt.savefig("GEOMETRIA.png",bbox_inches='tight')

def to_word(Z,I,M3,J_t,J_psi,flag_materiale,K,M,Tn,X,flag_lembo_vincolato):
    if flag_lembo_vincolato=="S":
        N=1+I
    else:
        N=I

    doc = docx.Document()
    #TITOLO
    p = doc.add_paragraph()
    doc.add_heading('TORSIONE COMPLETA', 0)
    #GEOMETRIA
    p = doc.add_paragraph()
    p.add_run(f"La trave ha una lunghezza pari a {Z[-1]} cm.")
    p = doc.add_paragraph()
    p.add_run(f"Sono stati applicati {I} momenti torcenti. I valori dei momenti torcenti sono stati riassunti nella seguente tabella.")
    table = doc.add_table(rows=1, cols=2)
    intestazioni = table.rows[0].cells
    intestazioni[0].text = 'ID Momento torcente'
    intestazioni[1].text = 'Momento torcente [kNcm]'
    j=1
    for M3i in M3:
        row_cells = table.add_row().cells
        row_cells[0].text = f"Mt{str(j)}"
        row_cells[1].text = str(M3i)
        j+=1
    p = doc.add_paragraph()
    p.add_run("Lo schema statico viene presentato nella seguente figura.")
    doc.add_picture('GEOMETRIA.png',width=Inches(6))

    p = doc.add_paragraph()
    p.add_run(f"La trave è realizzata in {materiale(flag_materiale)[2]}. Dunque E = {materiale(flag_materiale)[0]} MPa e G = {materiale(flag_materiale)[1]} MPa")
    
    p = doc.add_paragraph()
    p.add_run("La trave è caratterizzata dalle seguenti proprietà geometriche: ")
    table = doc.add_table(rows=1, cols=3)
    intestazioni = table.rows[0].cells
    intestazioni[0].text = 'ID tratto'
    intestazioni[1].text = 'Jt [cm^4]'
    intestazioni[2].text = 'Jpsi [cm^6]'
    
    for j in range(N):
        row_cells = table.add_row().cells
        row_cells[0].text = str(j+1)
        row_cells[1].text = str(J_t[j])
        row_cells[2].text = str(J_psi[j])

    p = doc.add_paragraph()
    p.add_run("Il coefficiente adimensionale k viene calcolato per ciascun tratto. ")

    table = doc.add_table(rows=1, cols=2)
    intestazioni = table.rows[0].cells
    intestazioni[0].text = 'ID tratto'
    intestazioni[1].text = 'k'
    
    for j in range(N):
        row_cells = table.add_row().cells
        row_cells[0].text = str(j+1)
        row_cells[1].text = str(K[j])  
    p.add_run("Si ricorda che per k>100 pressoche l'interezza del momento torcente è di tipo primario.")
    p = doc.add_paragraph()
    p.add_run("La matrice del sistema è : ")
    p = doc.add_paragraph()
    run=p.add_run(f"{M}")
    run.font.size = docx.shared.Pt(9)
    p = doc.add_paragraph()
    p.add_run("Il vettore dei termini noti è : ")
    p = doc.add_paragraph()
    p.add_run(f"{Tn}")
    p = doc.add_paragraph()
    p.add_run("Il vettore delle incognite è : ")
    p = doc.add_paragraph()
    p.add_run(f"{X}")
    p = doc.add_paragraph()
    p.add_run("L'andamento del momento torcente viene presentato nella seguente figura.")
    doc.add_picture("MOMENTO TORCENTE.png",width=Inches(6))
    p = doc.add_paragraph()
    p.add_run("Di seguito venogno forniti ulteriori grafici.")
    doc.add_picture("ALTRI GRAFICI.png",width=Inches(6))
    doc.save("RISULTATI.docx")

def risultati( I,K,Lj,J_t, J_psi, E,G , X):
    BIM=[]
    OMEGA=[]
    MTTOT=[]
    MTSEC=[]
    MTPRIM=[]
    OMEGA_P=[]
    p=1

    for j in range (I+1):
        
        z=range(0,int(Lj[j]),p)
        for zi in z:
            BIM.append(B(j,zi,K,Lj, J_psi, E,X)[0])
            OMEGA_P.append(Wp(zi,j,K ,Lj, X)[0])
            MTTOT.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[2][0])
            MTSEC.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[1][0])
            MTPRIM.append(Mt(zi,j,K,Lj,J_t, J_psi, E,G , X)[0][0])
            OMEGA.append(W(zi,j,K ,Lj, X)[0])

    z=[]
    for j in range (I+1):
        if j==0:
            start = int(0)
        else:
            start  = z[-1]+p
        
        zj=range(start,start+int(Lj[j]),p)
        
        z.extend(zj)

    return z, BIM, OMEGA_P, MTTOT, MTSEC, MTPRIM, OMEGA

def localize_floats(row):
    return [
        str(el).replace('.', ',') if isinstance(el, float) else el 
        for el in row
    ]

def esporta_csv(RES):
    dict_res ={ "z [cm]" : RES[0],
                "W [rad]" : RES[6],
                "W' [rad/cm]" : RES[2],
                "MT,tot [kNcm]" : RES[3],
                "MT,primario [kNcm]" : RES[5],
                "MT,secondario [kNcm]" : RES[4],
                "Bimomento [kNcm^2]" : RES[1]}
    

    df = pd.DataFrame(dict_res)

    df.to_csv('RISULTATI.csv', index = False, sep = ";", decimal= ",")


    
def main():

    print()
    print("______GEOMETRIA______")

    i=0
    while i <1:
        try:
            print()
            Ltot =float(input("Indicare la lunghezza della trave [m] : "))*100
            i=+1
        except:
            print()
            print("ATTENZIONE : Valore non valido")
            continue
    
    print()
    print("______VINCOLI______")
    
    i=0
    while i<1:
        try:
            print()
            flag_lembo_vincolato = input("L'estremo destro della trave è vincolato ? [S/N] : ")
            if flag_lembo_vincolato!="S" and flag_lembo_vincolato!="N":
                raise
            i+=1
        except:
            print()
            print("ATTENZIONE : Valore non valido.")

    if flag_lembo_vincolato == "S":
        print()
        input("La trave è vincolata ad entrambi gli estermi. Premi un tasto per continuare...")
    else:
        print()
        input("La trave è libera all'estremo destro. Premi un tasto per continuare...")

    print()
    print("______CARICHI______")
    print()

    I = int(input("Indicare il numero di momenti torcenti concentrati : "))
    M3=[]
    print()
    print("***INTENSITA' DEI MOMENTI TORCENTI***")
    i=0
    while i<=I-1:
        try:
            print()
            M3_i = float(input(f"Indicare l'intensità del momento torcente numero {i+1} [kNm]: "))*100
            i+=1
        except:
            print()
            print("ATTENZIONE : Valore non valido.")
            continue
        M3.append(M3_i)

    print()
    print("***POSIZIONE DEI MOMENTI TORCENTI***")
    print()
    print("ATTENZIONE: Non indicare la stessa posizione per due carichi differenti !!!")
    print()
    print("ATTENZIONE: Fare in modo che la posizione incrementi con l'indice del carico!!!")
    print()

    Z=[]
    i=0
    while i < I:
        try:
            print()
            Z_i = float(input(f"Indicare la posizione lungo la trave del momento torcente numero {i+1} [m]: "))*100
            i+=1
        except:
            print()
            print("ATTENZIONE : Valore non valido.")
            continue

        Z.append(Z_i)
        if max(Z) == Z_i:
            continue
        else:
            print(f"ERRORE : Il carico {i+1} ha una posizione non valida. L'esecuzione viene interrotta.")
            quit()
    print()
    print("ATTENZIONE : Indipendentemente dalla posizione indicata, in caso di lembo destro libero l'ultimo carico viene sempre posto in corrispondenza di tale estremità.")
    if flag_lembo_vincolato=="N":
        Z[-1]=Ltot

    Z.append(Ltot)
    Z.insert(0,0)

    if max(Z) > Ltot:
        print(f"ERRORE : Uno dei carichi è stato posto ad una distanza maggiore della lunghezza della trave. L'esecuzione viene interrotta.")
        quit()

    #LUNGHEZZE DEI VARI TRATTI IN CUI HO DIVISO LA TRAVE
    Lj=np.diff(Z)
    #plot_geometria( I,Lj,Z)

    

    print()
    print("______MATERIALE______")
    print()
    i=0
    while i<1:
        try:
            print()
            flag_materiale = int(input("Scegli materiale ( 1 - Acciaio, 2 - Calcestruzzo ) : "))
            if flag_materiale>2:
                raise
            i+=1
        except:
            print()
            print("ATTENZIONE : Valore non valido.")
            continue


    E=[]
    G=[]
    MAT=materiale(flag_materiale)
    for j in range(I+1):
        E_j=MAT[0]/10
        G_j=MAT[1]/10
        E.append(E_j)
        G.append(G_j)

    print()
    print("______INERZIE TORSIONALI______")
    print()

    INERZIE=Inerzie_torsionali(I)
    J_t = INERZIE[0]
    J_psi = INERZIE[1]

    print()
    print("______CALCOLO COEFFICIENTE K______")
    print()

    K=[]
    for i in range(I+1):
        k_j=  kj(Lj[i],J_t[i],J_psi[i],E[i],G[i])
        print (f" Tratto : {i+1}")
        print (f" Lunghezza [cm]: {Lj[i]}")
        print (f" Jt [cm^4]: {J_t[i]}")
        print (f" Jpsi [cm^6]: {J_psi[i]}")
        print (f" Mod. Young [kN/cm^2]: {E[i]}")
        print (f" Mod. Tang. [kN/cm^2]: {G[i]}")
        print (f" k : {float(k_j)}")
        K.append(k_j)
    

    if max(K)>100:
        print("Dato l'elevato valore del coefficiente k il software non riesce ad elaborare una soluzione : tutta la torsione è primaria e puo essere utilizzato SAP.")
        quit()
    else:
        print()
        print("______MATRICE DEL SISTEMA______")
        print()
        

        j=0
        if flag_lembo_vincolato=="S":
            M=np.zeros([4*(I+1),4*(I+1)])
            while j <(I+1):
                if j==0 :
                    print("Elemento : ")
                    print(f"{j} in {range(I)}")
                    print("---")
                    np.put( M[4*j,:] , [0,1,2,3] , cc_incastro_dsv(j,I,K,Lj,1) )
                    np.put( M[4*j+1,:] , [0,1,2,3] , cc_incastro_bimom(j,I,K,Lj,1)  )
                    np.put( M[4*j+2,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_dsv(j,I,K,Lj,2) )
                    np.put( M[4*j+3,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_bimom(j,I,K,Lj,2)  )
                    np.put( M[4*j+4,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_mom(j,I,K,Lj,J_t, J_psi, E,G, M3[j])  )
                    np.put( M[4*j+5,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_bimom(j,K,Lj, J_psi, E)  )
                    j+=1
            
                if j==I:
                        print("Elemento : ")
                        print(f"{j} in {range(I)}")
                        print("---")
                        np.put( M[4*j+2,:] , [4*I+0,4*I+1,4*I+2,4*I+3] , cc_incastro_dsv(j,I,K,Lj,1) )
                        np.put( M[4*j+3,:] , [4*I+0,4*I+1,4*I+2,4*I+3] , cc_incastro_bimom(j,I,K,Lj,1)  )
                        j+=1
                
                if j <I: 
                    print("Elemento : ")
                    print(f"{j} in {range(I)}")
                    print("---")
                    np.put( M[4*j+2,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_dsv(j,I,K,Lj,2) )
                    np.put( M[4*j+3,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_bimom(j,I,K,Lj,2)  )
                    np.put( M[4*j+4,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_mom(j,I,K,Lj,J_t, J_psi, E,G, M3[j])  )
                    np.put( M[4*j+5,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_bimom(j,K,Lj, J_psi, E)  )
                    j+=1
        
        else:
            M=np.zeros([4*(I),4*(I)])
            while j <(I):
                if j==I-1:
                    print("Elemento : ")
                    print(f"{j} in {range(I)}")
                    print("---")
                    np.put( M[4*j+2,:] , [4*I-4,4*I-3,4*I-2,4*I-1] , cc_lembo_libero_M(j,K,Lj,J_t, J_psi, E,G) )
                    np.put( M[4*j+3,:] , [4*I-4,4*I-3,4*I-2,4*I-1] , cc_lembo_libero_B(j,K,Lj, J_psi, E)  )
                    if I-1==0:
                        np.put( M[4*j,:] , [0,1,2,3] , cc_incastro_dsv(j,I,K,Lj,1) )
                        np.put( M[4*j+1,:] , [0,1,2,3] , cc_incastro_bimom(j,I,K,Lj,1)  )
                    j+=1
                if j==0 :
                    print("Elemento : ")
                    print(f"{j} in {range(I)}")
                    print("---")
                    np.put( M[4*j,:] , [0,1,2,3] , cc_incastro_dsv(j,I,K,Lj,1) )
                    np.put( M[4*j+1,:] , [0,1,2,3] , cc_incastro_bimom(j,I,K,Lj,1)  )
                    np.put( M[4*j+2,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_dsv(j,I,K,Lj,2) )
                    np.put( M[4*j+3,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_bimom(j,I,K,Lj,2)  )
                    np.put( M[4*j+4,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_mom(j,I,K,Lj,J_t, J_psi, E,G, M3[j])  )
                    np.put( M[4*j+5,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_bimom(j,K,Lj, J_psi, E)  )
                    j+=1
                if j <I-1:
                    print("Elemento : ")
                    print(f"{j} in {range(I)}")
                    print("---")
                    np.put( M[4*j+2,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_dsv(j,I,K,Lj,2) )
                    np.put( M[4*j+3,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_incastro_bimom(j,I,K,Lj,2)  )
                    np.put( M[4*j+4,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_mom(j,I,K,Lj,J_t, J_psi, E,G, M3[j])  )
                    np.put( M[4*j+5,:]  , [4*j+0,4*j+1,4*j+2,4*j+3,4*j+4,4*j+5,4*j+6,4*j+7 ] , cc_equ_bimom(j,K,Lj, J_psi, E)  )
                    j+=1


                    
        print((M))

        print()
        print("______VETTORE DEI TERMINI NOTI______")
        print()        
        
        j=0
        if flag_lembo_vincolato=="S":
            Tn = np.zeros([4*(I+1),1])
            for j in range (I):
                Tn[4*j+4]= M3[j]
        elif flag_lembo_vincolato=="N":
            Tn = np.zeros([4*(I),1])
            j=0
            while j<I:
                if j==I-1: 
                    Tn[-2] = M3[-1]
                    j+=1
                else:
                    Tn[4*j+4]= M3[j]
                    j+=1
            

        print(Tn)

        print()
        print("______SOLUZIONE DEL SISTEMA______")
        print() 
        C = np.linalg.inv(M)

        X= np.matmul(C,Tn)
        print(X)

        print()
        print("______RISULTATI______")
        print()
        plot_geometria( I,Lj,Z,flag_lembo_vincolato)
        grafici( I,K,Lj,J_t, J_psi, E,G , X,Z,flag_lembo_vincolato)
        to_word(Z,I,M3,J_t,J_psi,flag_materiale,K,M,Tn,X,flag_lembo_vincolato)
        RES= risultati( I,K,Lj,J_t, J_psi, E,G , X)
        esporta_csv(RES)


           

        
main()
